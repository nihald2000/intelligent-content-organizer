import logging
import tempfile
import os
from pathlib import Path
from typing import Optional, Dict, Any
import asyncio

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
from PIL import Image
import pytesseract

from .models import Document, DocumentType
import config

logger = logging.getLogger(__name__)

class DocumentParser:
    def __init__(self):
        self.config = config.config
    
    async def parse_document(self, file_path: str, filename: str) -> Document:
        """Parse a document and extract its content"""
        try:
            file_ext = Path(filename).suffix.lower()
            file_size = os.path.getsize(file_path)
            
            # Determine document type and parse accordingly
            if file_ext == '.pdf':
                content = await self._parse_pdf(file_path)
                doc_type = DocumentType.PDF
            elif file_ext == '.txt':
                content = await self._parse_text(file_path)
                doc_type = DocumentType.TEXT
            elif file_ext == '.docx':
                content = await self._parse_docx(file_path)
                doc_type = DocumentType.DOCX
            elif file_ext in ['.png', '.jpg', '.jpeg', '.bmp', '.tiff']:
                content = await self._parse_image(file_path)
                doc_type = DocumentType.IMAGE
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            # Create document object
            document = Document(
                id=self._generate_document_id(),
                filename=filename,
                content=content,
                doc_type=doc_type,
                file_size=file_size,
                metadata={
                    "file_extension": file_ext,
                    "content_length": len(content),
                    "word_count": len(content.split()) if content else 0
                }
            )
            
            logger.info(f"Successfully parsed document: {filename}")
            return document
            
        except Exception as e:
            logger.error(f"Error parsing document {filename}: {str(e)}")
            raise
    
    async def _parse_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            content += f"\n--- Page {page_num + 1} ---\n"
                            content += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {str(e)}")
                        continue
            
            return content.strip()
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise
    
    async def _parse_text(self, file_path: str) -> str:
        """Read plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            return content.strip()
        except Exception as e:
            logger.error(f"Error parsing text file: {str(e)}")
            raise
    
    async def _parse_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(file_path)
            content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content += " | ".join(row_text) + "\n"
            
            return content.strip()
        except Exception as e:
            logger.error(f"Error parsing DOCX file: {str(e)}")
            raise
    
    async def _parse_image(self, file_path: str) -> str:
        """Extract text from image using OCR"""
        try:
            # First try with OCR service if available
            if hasattr(self, 'ocr_service') and self.ocr_service:
                logger.info(f"Using OCR service for image: {file_path}")
                text = await self.ocr_service.extract_text_from_image(file_path)
                if text:
                    return text
            
            # Fallback to direct pytesseract
            logger.info(f"Using direct pytesseract for image: {file_path}")
            image = Image.open(file_path)
            
            # Perform OCR
            content = pytesseract.image_to_string(
                image,
                lang=self.config.OCR_LANGUAGE,
                config='--psm 6'  # Assume a single uniform block of text
            )
            
            return content.strip()
        except Exception as e:
            logger.error(f"Error performing OCR on image: {str(e)}")
            # Return empty string if OCR fails
            return ""
    
    def _generate_document_id(self) -> str:
        """Generate a unique document ID"""
        import uuid
        return str(uuid.uuid4())
    
    async def extract_metadata(self, file_path: str, content: str) -> Dict[str, Any]:
        """Extract additional metadata from the document"""
        try:
            metadata = {}
            
            # Basic statistics
            metadata["content_length"] = len(content)
            metadata["word_count"] = len(content.split()) if content else 0
            metadata["line_count"] = len(content.splitlines()) if content else 0
            
            # File information
            file_stat = os.stat(file_path)
            metadata["file_size"] = file_stat.st_size
            metadata["created_time"] = file_stat.st_ctime
            metadata["modified_time"] = file_stat.st_mtime
            
            # Content analysis
            if content:
                # Language detection (simple heuristic)
                metadata["estimated_language"] = self._detect_language(content)
                
                # Reading time estimation (average 200 words per minute)
                metadata["estimated_reading_time_minutes"] = max(1, metadata["word_count"] // 200)
            
            return metadata
        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            return {}
    
    def _detect_language(self, content: str) -> str:
        """Simple language detection based on character patterns"""
        # This is a very basic implementation
        # In production, you might want to use a proper language detection library
        if not content:
            return "unknown"
        
        # Count common English words
        english_words = ["the", "and", "or", "but", "in", "on", "at", "to", "for", "of", "with", "by", "from", "as", "is", "was", "are", "were", "be", "been", "have", "has", "had", "do", "does", "did", "will", "would", "could", "should", "may", "might", "can", "this", "that", "these", "those"]
        
        words = content.lower().split()
        english_count = sum(1 for word in words if word in english_words)
        
        if len(words) > 0 and english_count / len(words) > 0.1:
            return "en"
        else:
            return "unknown"